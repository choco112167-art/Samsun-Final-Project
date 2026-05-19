"""팩트체크 파이프라인 가이드 문서 생성 (v2.0 최신화)"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)


def sf(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def title(text, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    sf(r, size, bold)
    return p


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, 13, bold=True, color=(31, 73, 125))
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, 11, bold=True, color=(68, 114, 196))
    return p


def body(text, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, size)
    return p


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, 8.5, italic=True, color=(128, 128, 128))
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0, 70, 0)
    return p


def tbl(headers, rows, col_widths, header_color='1F497D'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        sf(r, 9, bold=True)
        r.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = etree.SubElement(tcPr, f'{{{W}}}shd')
        shd.set(f'{{{W}}}val', 'clear')
        shd.set(f'{{{W}}}color', 'auto')
        shd.set(f'{{{W}}}fill', header_color)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            sf(r, 9)
    for row_obj in t.rows:
        for j, cell in enumerate(row_obj.cells):
            cell.width = Cm(col_widths[j])
    return t


def divider():
    p = doc.add_paragraph('─' * 80)
    p.paragraphs[0] if hasattr(p, 'paragraphs') else None
    r = p.runs[0] if p.runs else p.add_run()
    sf(r, 7, color=(180, 180, 180))
    return p


# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
title('삼선뉴스 팩트체커 설계 문서')
title('언론사 신뢰도 점수 근거  ·  팩트체크 파이프라인 상세', size=12, bold=False)
doc.add_paragraph()
title('작성일: 2026-04-22    최종 수정: 2026-05-19    담당: 이동우    버전: v2.0', size=9, bold=False)
doc.add_paragraph()

note('v2.0 변경사항: ① INSIGHT 라벨 추가 (2026-04-30) '
     '② CoVe 소스 신뢰도 보정 적용 (2026-05-19) '
     '③ Reddit source_type 오기재 방어 코드 추가 (2026-05-19)')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 1. 전체 파이프라인 위치
# ══════════════════════════════════════════════════════════════
h1('1. 전체 파이프라인에서 팩트체커 위치')
body('팩트체커는 RSS 수집 직후, 번역·요약 전에 위치합니다. 쓰레기 기사(칼럼·루머·커뮤니티 소문)를 '
     'AI가 번역하면 사실인 것처럼 포장되므로, 앞 단계에서 선제적으로 걸러냅니다.')
doc.add_paragraph()

tbl(
    ['단계', '담당', '출력'],
    [
        ['RSS 수집', '이상준', 'Article 객체 (title, content, source, source_type)'],
        ['[ 팩트체커 ]  ← 여기', '이동우', 'fact_label (FACT/RUMOR/UNVERIFIED/INSIGHT/DROP)'],
        ['번역·요약', '이동우', 'translation, summary_formal, summary_casual'],
        ['임베딩 + RAG 추천', '강주찬', 'embedding VECTOR(1024), pgvector'],
        ['FastAPI → React 화면', '강주찬·정수민', '개인화 피드 서빙'],
    ],
    col_widths=[4, 2.5, 9]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 2. 5가지 라벨
# ══════════════════════════════════════════════════════════════
h1('2. 5가지 라벨 — 정확한 판정 기준')
note('v2.0 추가: INSIGHT 라벨 (v1.0 대비 변경)')
doc.add_paragraph()

tbl(
    ['라벨', '의미', 'DB 저장', '화면 표시'],
    [
        ['FACT',       '사실 확인됨',                    'O', '정상 서빙'],
        ['RUMOR',      '루머/미확인 — 출처는 있지만 신뢰 불가', 'O', '루머 주의 뱃지'],
        ['UNVERIFIED', '검증 불가 — human-in-the-loop 필요', 'O', '검토 중 플래그'],
        ['INSIGHT',    '전문가 사설/분석 — TIER 0-1 Opinion', 'O', '분석글 뱃지 (v2.0 추가)'],
        ['DROP',       '노이즈/저품질 — 커뮤니티 잡글·칼럼', 'X', '서빙 안 함'],
    ],
    col_widths=[2.5, 6, 2, 5]
)
doc.add_paragraph()

h2('2-1. DROP — 처음부터 받지 않는다')
body('아래 두 조건 중 하나 충족 시 즉시 DROP. DB 저장 없음.')
body('  조건 A: 출처 자체가 DROP 등급 (COMMUNITY_NOISE 티어 — Reddit 잡글, 익명 커뮤니티)')
body('  조건 B: 언론사 기사이지만 Opinion 패턴이 2개 이상 탐지됨')
body('    → opinion, editorial, i think, in my view, the case for, frankly 등')
body('    → 1개는 오탐 방지로 무시, 2개 이상일 때만 DROP')
doc.add_paragraph()

h2('2-2. INSIGHT — 전문가 사설은 버리지 않는다  [v2.0 신규]')
body('TIER 0(학술/기관) 또는 TIER 1(공식 언론사)에서 Opinion 패턴 2개+ 탐지 시.')
body('v1.0에서는 이 경우도 DROP으로 처리했으나, 전문가 사설은 콘텐츠 가치가 있어 INSIGHT로 저장.')
body('  → MIT Tech Review 칼럼, IEEE 기고 등이 해당')
body('  → conf=0.85 고정, DB 저장 O, 화면에 "분석글" 뱃지 표시')
doc.add_paragraph()

h2('2-3. RUMOR — 출처는 있지만 신뢰 불가')
body('아래 중 하나라도 해당 시 RUMOR. DB 저장, 루머 주의 뱃지 표시.')
body('  Step 1: 강한 루머 신호 패턴 탐지')
body('    → allegedly, reportedly, purportedly, debunked, misinformation,')
body('       가짜 뉴스, 허위, 루머, 논란 등 (conf=0.80)')
body('  Step 2: Google Fact Check API → RUMOR 판정 반환 (conf=0.90)')
body('  Step 3A: Gemini Advisor → confidence ≥ 0.80으로 RUMOR 판정')
doc.add_paragraph()

h2('2-4. UNVERIFIED — 사람이 직접 검토해야 한다')
body('아래 중 하나에 해당 시 UNVERIFIED. DB 저장, human-in-the-loop 플래그.')
body('  Step 1: 약한 루머 신호만 있음 (may be, suggests, possibly, expected to 등)')
body('  Step 1: CREDIBLE_LEAK 출처 (exclusive, sources say, 단독 보도)')
body('  Step 3A: Gemini confidence 0.50~0.79 (확신 부족)')
body('  Step 3B CoVe: 독립 검증 5문 중 일부만 통과 (agreement_ratio 낮음)')
body('  Step 3C DebateCV: 검사 vs 변호인 vs 판사 토론 후 2:1 판정')
doc.add_paragraph()

h2('2-5. FACT — 믿을 수 있다')
body('아래 중 하나에 해당 시 FACT 처리.')
body('  Step 1 자동: 공식 언론사 + 루머 신호 없음 → FACT_AUTO (Gemini 호출 없이 처리)')
body('  Step 2: Google Fact Check API → FACT 판정 (conf=0.90)')
body('  Step 3A: Gemini confidence ≥ 0.80으로 FACT 판정')
body('  Step 3B CoVe: 소스 신뢰도 보정 후 confidence ≥ 0.80  [v2.0 수정]')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 3. Step 0~3C 전체 흐름
# ══════════════════════════════════════════════════════════════
h1('3. Step 0~3C 전체 흐름')
doc.add_paragraph()

tbl(
    ['Step', '모듈', '동작', '출력'],
    [
        ['Step 0', 'channel_config.py',
         'source → ChannelProfile 조회\nCOMMUNITY_NOISE / MEDIA_OPINION 티어 → DROP',
         'DROP (conf=1.0) 또는 계속'],
        ['Step 1', 'signal_detector.py',
         'Opinion 패턴 2개+ → DROP or INSIGHT(TIER0-1)\n'
         'STRONG 루머 패턴 → RUMOR (conf=0.80)\n'
         '신호 없음 + 공식미디어 → FACT_AUTO\n'
         'CREDIBLE_LEAK / COMMUNITY → NEEDS_VERIFICATION',
         'INSIGHT / RUMOR / FACT / NEEDS_VERIFICATION'],
        ['Step 2', 'google_fc_api.py',
         'ClaimReview DB 200+ 기관 조회\nAI 뉴스 매칭률 ~15.8%',
         '매칭 시 FACT/RUMOR (conf=0.90)'],
        ['Step 3A', 'gemini_advisor.py',
         'importance_score 산출\nPass A: 문체 분석 (감정어·헤징 밀도)\nPass B: Google Search Grounding + 상식 추론',
         'conf ≥ 0.80 → 종료\nconf < 0.80 → 3B 진입'],
        ['Step 3B', 'cove_verifier.py',
         '5개 독립 검증 질문 (존재/수치/출처/맥락/신규성)\n'
         '소스 신뢰도 보정: TIER 0-1은 credibility×0.15 추가  [v2.0]\n'
         'conf ≥ 0.80 or importance < 0.70 → 종료',
         'FACT/RUMOR/UNVERIFIED\n+ conf 보정값'],
        ['Step 3C', 'debate_agents.py',
         'importance ≥ 0.70 AND 여전히 불확실 시 발동\n'
         'Prosecutor(반론) vs Defender(지지) vs Judge(판정)\n'
         '만장일치 conf ≥ 0.90 / 2:1 conf 0.70~0.80',
         'FACT/RUMOR/UNVERIFIED\n+ debate_outcome'],
    ],
    col_widths=[1.8, 3.5, 7.5, 3]
)
doc.add_paragraph()

h2('importance_score 산출 공식')
code_block(
    'importance_score = (\n'
    '    0.40 * min(model_name_count / 3, 1.0)  # GPT-5, Gemini 3, Claude 4 등 모델명 수\n'
    '  + 0.30 * has_benchmark_number             # MMLU 89.3%, SWE-bench 64.3% 등\n'
    '  + 0.20 * has_superlative_claim            # best, SOTA, world record, unprecedented\n'
    '  + 0.10 * is_breaking_news                 # breaking, exclusive, leaked (제목 한정)\n'
    ')  # 0.0~1.0, ≥ 0.70이면 Step 3C DebateCV 발동'
)
doc.add_paragraph()

h2('CoVe 소스 신뢰도 보정  [v2.0 신규]')
body('문제: CoVe LLM이 AI 기술 기사의 수치 검증에 실패하면 confident=False → UNVERIFIED 과다 판정.')
body('해결: TIER 0(학술/기관) / TIER 1(공식 미디어) 소스에 한해 CoVe confidence에 보정값 추가.')
code_block(
    '# pipeline.py — Step 3B 결과 처리 (v2.0 수정)\n'
    'if tier in ("ACADEMIC_INSTITUTIONAL", "MEDIA_OFFICIAL"):\n'
    '    credibility_boost = profile.credibility_score * 0.15\n'
    '    cove_conf_adjusted = min(cove_result.confidence + credibility_boost, 1.0)\n'
    'else:\n'
    '    cove_conf_adjusted = cove_result.confidence\n'
    '\n'
    '# 예시: TechCrunch (credibility=0.82) → +0.12 보정\n'
    '#       MIT Tech Review (credibility=0.95) → +0.14 보정\n'
    '# 근거: Baly et al. EMNLP 2018 — 소스 신뢰도가 팩트체크의 강력한 사전 확률'
)
doc.add_paragraph()

h2('Reddit source_type 방어 코드  [v2.0 신규]')
body('문제: testset CSV에서 Reddit 기사 일부가 source_type=media로 잘못 기재 → FACT_AUTO 오분류.')
code_block(
    '# pipeline.py — Step 0 직후 (v2.0 추가)\n'
    'if "reddit" in source.lower() and source_type == "media":\n'
    '    source_type = "community"  # 소스명 기반 강제 보정'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 4. 시나리오 4가지
# ══════════════════════════════════════════════════════════════
h1('4. 가짜 기사 데모 — 5가지 시나리오')
note('v2.0 추가: 시나리오 5 (INSIGHT)')
doc.add_paragraph()

scenarios = [
    ('시나리오 1 — DROP (칼럼 기사)', [
        ('제목', '"Why I Think Generative AI Will Destroy Creativity — An Opinion"'),
        ('출처', 'TechCrunch (media)'),
        ('판정 경로',
         'Step 0: TechCrunch → MEDIA_OFFICIAL 티어, DROP 아님\n'
         'Step 1 Opinion 탐지: in my view / i believe / frankly / the case against / i think → 5개 (≥2개)'),
        ('결과', 'DROP (conf=0.90) — DB 저장 안 함'),
    ]),
    ('시나리오 2 — RUMOR (미확인 인수설)', [
        ('제목', '"OpenAI Allegedly Planning to Acquire Anthropic, Sources Say"'),
        ('출처', 'VentureBeat AI (media)'),
        ('판정 경로',
         'Step 0: VentureBeat → MEDIA_CREDIBLE_LEAK 티어\n'
         'Step 1 STRONG 탐지: allegedly / reportedly / unverified / debunked → 4개'),
        ('결과', 'RUMOR (conf=0.80) — 루머 주의 뱃지'),
    ]),
    ('시나리오 3 — UNVERIFIED (소식통 단독보도)', [
        ('제목', '"Google May Be Planning to Release Gemini Ultra 2 Next Quarter"'),
        ('출처', 'The Decoder (media)'),
        ('판정 경로',
         'Step 0: MEDIA_OFFICIAL 티어\n'
         'Step 1: exclusive / sources familiar with → CREDIBLE_LEAK, FACT_AUTO 불가\n'
         '        may be / expected to / believed to → NEEDS_VERIFICATION\n'
         'Step 2: Google FC API 미매칭\n'
         'Step 3A: Gemini conf=0.68 → Step 3B 진입\n'
         'Step 3B: 5문 중 3문 검증 불가 → conf=0.62 + 보정(0.82×0.15=+0.12) = 0.74\n'
         '         importance=0.70 → Step 3C 진입\n'
         'Step 3C: DebateCV 2:1 → UNVERIFIED (conf=0.75)'),
        ('결과', 'UNVERIFIED — 팀원 수동 검토 필요'),
    ]),
    ('시나리오 4 — FACT (공식 발표)', [
        ('제목', '"Anthropic Releases Claude 4 with 200K Context Window"'),
        ('출처', 'MIT Technology Review (media)'),
        ('판정 경로',
         'Step 0: MIT Tech Review → ACADEMIC_INSTITUTIONAL 티어 (credibility=0.95)\n'
         'Step 1: Opinion 0개 / Credible Leak 0개 / 루머 신호 0개\n'
         '        ACADEMIC_INSTITUTIONAL + 루머 신호 없음 → FACT_AUTO'),
        ('결과', 'FACT (conf=0.95) — Gemini 호출 없이 자동 처리'),
    ]),
    ('시나리오 5 — INSIGHT (전문가 사설)  [v2.0 신규]', [
        ('제목', '"The Case Against AI Regulation: Why Government Oversight Will Stifle Innovation"'),
        ('출처', 'MIT Technology Review (media)'),
        ('판정 경로',
         'Step 0: MIT Tech Review → ACADEMIC_INSTITUTIONAL 티어 (TIER 0)\n'
         'Step 1 Opinion 탐지: the case against / i think / i believe / in my view → 4개 (≥2개)\n'
         '        → TIER 0 소스이므로 DROP 대신 INSIGHT 처리 (v1.0은 DROP)'),
        ('결과', 'INSIGHT (conf=0.85) — DB 저장 O, 분석글 뱃지 표시'),
    ]),
]

for title_text, items in scenarios:
    h2(title_text)
    for k, v in items:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{k}: ')
        sf(r1, 9, bold=True)
        r2 = p.add_run(v)
        sf(r2, 9)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 5. 언론사 신뢰도 점수
# ══════════════════════════════════════════════════════════════
h1('5. 7개 언론사 신뢰도 점수 근거')
doc.add_paragraph()

tbl(
    ['언론사', '티어', 'MBFC 사실보도', 'MBFC 편향', 'NewsGuard', '우리 점수'],
    [
        ['MIT Technology Review', 'TIER 0', 'VERY HIGH', 'Pro-Science', '녹색(75+)', '0.95'],
        ['IEEE Spectrum',         'TIER 0', 'VERY HIGH', 'Pro-Science', '녹색(75+)', '0.93'],
        ['The Guardian Tech',     'TIER 1', 'HIGH',      'Left-Center', '녹색(75+)', '0.88'],
        ['TechCrunch',            'TIER 1', 'HIGH',      'Left-Center', '녹색(75+)', '0.82'],
        ['The Decoder',           'TIER 1', '미평가',     '미평가',       '미평가',    '0.82'],
        ['The Verge',             'TIER 1', 'HIGH',      'Left-Center', '녹색(75+)', '0.80'],
        ['VentureBeat AI',        'TIER 2', 'HIGH(비공식)', 'Center',   '미평가',    '0.72'],
    ],
    col_widths=[4, 2, 2.5, 2.5, 2, 2]
)
doc.add_paragraph()

h2('4팩터 신뢰도 공식')
code_block(
    'credibility_score = (\n'
    '    0.40 * source_tier_score      # TIER 기본값 (ACADEMIC=0.95, MEDIA_OFF=0.85 등)\n'
    '  + 0.25 * track_record_score     # 과거 팩트체크 실패율 역수 (NewsGuard DB)\n'
    '  + 0.20 * transparency_score     # 운영자·출처 공개 여부\n'
    '  + 0.15 * citation_score         # 원문·논문 링크 포함 비율\n'
    ')\n'
    '# 근거: Baly et al. EMNLP 2018 — 1,000개 뉴스 소스 신뢰도 DB'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 6. 소스 유형별 티어
# ══════════════════════════════════════════════════════════════
h1('6. 소스 유형별 티어 체계 (7 티어)')
doc.add_paragraph()

tbl(
    ['티어', '소스 유형', '예시', '기본 점수', 'FACT_AUTO', '시작 Step'],
    [
        ['TIER 0', '학술/기관',        'MIT TR, IEEE Spectrum, arXiv', '0.90~1.00', 'O', 'Step 1'],
        ['TIER 1', '공식 언론사',      'TechCrunch, Verge, Guardian, Decoder', '0.78~0.92', 'O', 'Step 1'],
        ['TIER 2', '전문미디어/유출',  'VentureBeat, Wired', '0.60~0.78', 'X (CREDIBLE_LEAK)', 'Step 2'],
        ['TIER 3', '인플루언서·유튜브','Two Minute Papers, Karpathy', '0.30~0.60', 'X', 'Step 3A'],
        ['TIER 4', '커뮤니티',         'Reddit r/ML, r/LocalLLaMA', '0.20~0.45', 'X', 'Step 2'],
        ['TIER 5', '텔레그램·익명',    'AI News 채널, LLM Leaks', '0.05~0.25', 'X', 'RUMOR 고정'],
        ['DROP',   'COMMUNITY_NOISE / MEDIA_OPINION', 'Reddit r/artificial, 칼럼', '0.00', '-', 'Step 0 종료'],
    ],
    col_widths=[2, 3, 5, 2.5, 3, 2.5]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 7. 크로스 소스 부스트
# ══════════════════════════════════════════════════════════════
h1('7. 크로스 소스 검증 (cross_source_boost)')
body('동일 claim을 TIER 1 언론사 복수가 독립 보도 시 confidence 상향 조정.')
code_block(
    'def cross_source_boost(claim, all_sources) -> float:\n'
    '    tier1_count = count_tier1_coverage(claim, all_sources)\n'
    '    if tier1_count >= 3:  return +0.20  # TechCrunch + Guardian + Verge\n'
    '    elif tier1_count == 2: return +0.10\n'
    '    elif tier1_count == 1: return +0.05\n'
    '    return 0.0\n'
    '# 텔레그램 루머 → 24h 내 TIER 1 2개+ 독립 보도 → UNVERIFIED → 승격 가능'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 8. UNVERIFIED 수동 검토 기준
# ══════════════════════════════════════════════════════════════
h1('8. UNVERIFIED → 팀원 수동 검토 기준')
tbl(
    ['상황', '액션'],
    [
        ['CoVe confidence 0.50~0.79',        '기사 직접 읽고 FACT/RUMOR 오버라이드'],
        ['DebateCV 2:1 판정',                 'reasoning_trace 읽고 판사 논리 동의 여부 확인'],
        ['CREDIBLE_LEAK 소식통 인용 기사',    '원 소식통 신뢰도 직접 판단'],
        ['Gemini 2-Pass 실패 (API 오류)',     '자동 UNVERIFIED 처리 → 수동 검토 필수'],
        ['생의학·비AI 도메인 기사 UNVERIFIED', '도메인 밖 — CoVe 검증 불가, 수동 판정'],
    ],
    col_widths=[6, 9.5]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 9. 평가 결과 (v2.0 기준)
# ══════════════════════════════════════════════════════════════
h1('9. 파이프라인 평가 결과 (2026-05-19, 201건)')
tbl(
    ['지표', '목표', '초기 결과', '최종 결과 (v2.0)', '달성'],
    [
        ['신뢰도 분류 정확도', '≥ 80%', '90.5% (182/201)', '98.5% (198/201)', '✅'],
        ['RUMOR recall',     '≥ 0.75', '1.000 (11/11)',    '1.000 (11/11)',   '✅'],
        ['FACT F1',          '참고용',  '0.928',            '0.989',           '✅'],
        ['처리 속도',         '≤ 5초',  '1.46초/건',        '1.46초/건',       '✅'],
        ['DROP recall',      '참고용',  '1.000 (20/20)',    '1.000 (20/20)',   '✅'],
    ],
    col_widths=[4, 2, 3.5, 3.5, 1.5]
)
doc.add_paragraph()
note('v2.0 CoVe 소스 신뢰도 보정으로 FACT→UNVERIFIED 오분류 17건 중 16건 정정. '
     '잔여 3건: 생의학 도메인 1건(CoVe 검증 불가) + Reddit source_type 오기재 2건.')

# ══════════════════════════════════════════════════════════════
# 푸터
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
title('삼선뉴스 프로젝트 — 내부 공유용 문서  |  v2.0  |  2026-05-19', size=8, bold=False)

out = r'C:\Users\이동우\Desktop\factcheck_pipeline_guide_v2.docx'
doc.save(out)
print(f'저장 완료: {out}')
